from datetime import datetime
from rich.console import Console
from rich.panel import Panel
from rich.table import Table
from rich.text import Text
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

console = Console()

STATUS_COLORS = {
    "ok": "green",
    "warning": "yellow",
    "critical": "red",
}


def _status_text(status: str) -> Text:
    color = STATUS_COLORS.get(status.lower(), "white")
    return Text(status.upper(), style=f"bold {color}")


def print_terminal_report(summary: dict, analysis: dict) -> None:
    date_range = summary.get("date_range", {})
    start = date_range.get("start", "N/A")
    end = date_range.get("end", "N/A")

    console.print(
        Panel(
            f"[bold]Operations Report[/bold]\nPeriod: {start} to {end}  |  Rows: {summary.get('row_count', 'N/A')}  |  Defect Rate: {summary.get('defect_rate', 'N/A')}%",
            title="[bold blue]Ops Report Generator[/bold blue]",
            border_style="blue",
        )
    )

    # KPI table
    kpi_table = Table(title="Key Performance Indicators", show_header=True, header_style="bold cyan")
    kpi_table.add_column("KPI", style="bold")
    kpi_table.add_column("Value")
    kpi_table.add_column("Status")

    for kpi in analysis.get("kpis", []):
        kpi_table.add_row(kpi["name"], kpi["value"], _status_text(kpi["status"]))

    console.print(kpi_table)

    # Anomalies table
    anomaly_table = Table(title="Anomalies Detected", show_header=True, header_style="bold magenta")
    anomaly_table.add_column("Description", no_wrap=False, max_width=70)
    anomaly_table.add_column("Severity")

    for anomaly in analysis.get("anomalies", []):
        anomaly_table.add_row(anomaly["description"], _status_text(anomaly["severity"]))

    console.print(anomaly_table)

    # Recommended actions
    console.print("\n[bold cyan]Recommended Actions[/bold cyan]")
    for i, action in enumerate(analysis.get("actions", []), 1):
        console.print(f"  [bold]{i}.[/bold] {action}")

    console.print()


def generate_word_report(summary: dict, analysis: dict, output_path: str) -> str:
    doc = Document()

    date_range = summary.get("date_range", {})
    start = date_range.get("start", "N/A")
    end = date_range.get("end", "N/A")

    # Title
    title = doc.add_heading(f"Operations Report: {start} to {end}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    generated_at = datetime.now().strftime("%B %d, %Y at %H:%M")
    meta = doc.add_paragraph(
        f"Generated on {generated_at}  |  Rows analyzed: {summary.get('row_count', 'N/A')}  |  Overall defect rate: {summary.get('defect_rate', 'N/A')}%"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # KPIs section
    doc.add_heading("Key Performance Indicators", level=1)
    kpi_table = doc.add_table(rows=1, cols=3)
    kpi_table.style = "Table Grid"
    hdr_cells = kpi_table.rows[0].cells
    hdr_cells[0].text = "KPI"
    hdr_cells[1].text = "Value"
    hdr_cells[2].text = "Status"
    for cell in hdr_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    for kpi in analysis.get("kpis", []):
        row_cells = kpi_table.add_row().cells
        row_cells[0].text = kpi["name"]
        row_cells[1].text = kpi["value"]
        status = kpi["status"].lower()
        run = row_cells[2].paragraphs[0].add_run(status.upper())
        run.bold = True
        if status == "critical":
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        elif status == "warning":
            run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
        else:
            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

    doc.add_paragraph()

    # Anomalies section
    doc.add_heading("Anomalies Detected", level=1)
    anomaly_table = doc.add_table(rows=1, cols=2)
    anomaly_table.style = "Table Grid"
    hdr_cells = anomaly_table.rows[0].cells
    hdr_cells[0].text = "Description"
    hdr_cells[1].text = "Severity"
    for cell in hdr_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    for anomaly in analysis.get("anomalies", []):
        row_cells = anomaly_table.add_row().cells
        row_cells[0].text = anomaly["description"]
        severity = anomaly["severity"].lower()
        run = row_cells[1].paragraphs[0].add_run(severity.upper())
        run.bold = True
        if severity == "critical":
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        elif severity == "warning":
            run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
        else:
            run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

    doc.add_paragraph()

    # Recommended actions section
    doc.add_heading("Recommended Actions", level=1)
    for i, action in enumerate(analysis.get("actions", []), 1):
        doc.add_paragraph(f"{i}. {action}", style="List Number")

    doc.add_paragraph()

    # Footer
    footer_para = doc.add_paragraph("Generated by ops_report_generator")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(output_path)
    return output_path


if __name__ == "__main__":
    sample_summary = {
        "row_count": 32,
        "date_range": {"start": "2024-01-01", "end": "2024-01-11"},
        "stats": {
            "units_produced": {"mean": 402.59, "min": 0.0, "max": 465.0, "null_count": 0},
            "defects": {"mean": 13.19, "min": 0.0, "max": 150.0, "null_count": 0},
            "downtime_hours": {"mean": 0.69, "min": 0.0, "max": 9.5, "null_count": 0},
        },
        "defect_rate": 3.28,
        "lines": ["L1", "L2", "L3"],
        "shifts": ["afternoon", "morning", "night"],
    }

    sample_analysis = {
        "kpis": [
            {"name": "Defect Rate", "value": "3.28%", "status": "warning"},
            {"name": "Avg Units Produced", "value": "402.59", "status": "ok"},
            {"name": "Avg Downtime", "value": "0.69 hrs", "status": "ok"},
            {"name": "Max Single-Day Defects", "value": "150", "status": "critical"},
            {"name": "Max Downtime Event", "value": "9.5 hrs", "status": "critical"},
        ],
        "anomalies": [
            {
                "description": "Line L1 on 2024-01-05 reported 150 defects — 12x above average.",
                "severity": "critical",
            },
            {
                "description": "Line L2 on 2024-01-06 recorded 9.5 hours of downtime with zero production.",
                "severity": "critical",
            },
            {
                "description": "Three shifts reported zero units with no downtime — possible missing data.",
                "severity": "warning",
            },
        ],
        "actions": [
            "Investigate root cause of L1's 150-defect spike on Jan 5.",
            "Review L2's 9.5-hour downtime on Jan 6.",
            "Audit idle shifts with no production logged.",
            "Set automated alerts for defect counts exceeding 30 per shift.",
        ],
    }

    print_terminal_report(sample_summary, sample_analysis)
    path = generate_word_report(sample_summary, sample_analysis, "test_report.docx")
    print(f"Word report saved to: {path}")
